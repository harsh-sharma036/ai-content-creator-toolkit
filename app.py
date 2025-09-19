# import streamlit as st
# import google.generativeai as genai
# import requests
# import os

# from dotenv import load_dotenv
# from pytrends.request import TrendReq
# import pandas as pd
# from docx import Document
# from io import BytesIO

# load_dotenv()

# st.set_page_config(
#     page_title="AI Content Generator",
#     page_icon="✨",
#     layout="centered",
#     initial_sidebar_state="expanded"
# )

# st.markdown("""
#     <style>
#         body, .stApp {
#             background-color: #181824;
#             color: #f1f1f1;
#         }
#         .stButton>button {
#             background-color: #ff4b4b;
#             color: white;
#             border-radius: 8px;
#             font-weight: bold;
#         }
#         .stTextInput>div>input, .stTextArea>div>textarea, .stSelectbox>div>div>div>div {
#             background-color: #232339;
#             color: #f1f1f1;
#             border-radius: 6px;
#         }
#         .stSidebar {
#             background-color: #232339;
#         }
#         .stRadio > div {
#             gap: 0.5rem;
#         }
#         .stMarkdown a {
#             color: #ff4b4b;
#         }
#     </style>
# """, unsafe_allow_html=True)

# st.sidebar.markdown("""
# <div style="text-align:center">
#     <h2 style="color:#ff4b4b;">Creator Toolkit</h2>
#     <p style="font-size:15px;">Your all-in-one AI-powered content lab.<br>Generate, research, and optimize for every platform!</p>
#     <hr>
# </div>
# """, unsafe_allow_html=True)

# st.sidebar.title("🛠️ Features")
# feature = st.sidebar.radio(
#     "Select a feature:",
#     [
#         "Content Generation",
#         "Niche & Keyword Research",
#         "Meme Generation",
#         "YouTube SEO Metadata",
#         "Caption-to-Image Prompt"
#     ]
# )

# gemini_api_key = os.getenv("GEMINI_API_KEY")
# if not gemini_api_key:
#     st.error("Please set your GEMINI_API_KEY environment variable to use AI features. For example, export GEMINI_API_KEY=your_key_here")

# def get_gemini_model():
#     if gemini_api_key:
#         genai.configure(api_key=gemini_api_key)
#         return genai.GenerativeModel('gemini-1.5-flash')
#     return None

# def get_trending_keywords(query, gprop='youtube'):
#     try:
#         pytrends = TrendReq(hl='en-US', tz=360)
#         pytrends.build_payload([query], cat=0, timeframe='now 7-d', geo='', gprop=gprop)
#         related = pytrends.related_queries()
#         if query in related and related[query]['top'] is not None:
#             return related[query]['top']['query'].tolist()[:10]
#     except Exception:
#         pass
#     return []

# @st.cache_data
# def get_meme_templates():
#     response = requests.get("https://api.memegen.link/templates/")
#     if response.status_code == 200:
#         return response.json()
#     else:
#         st.error("Failed to fetch meme templates from Memegen.link.")
#         return []

# def export_to_docx(text, filename="output.docx"):
#     doc = Document()
#     for line in text.split('\n'):
#         doc.add_paragraph(line)
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     st.download_button(
#         label="Export to DOCX",
#         data=buffer,
#         file_name=filename,
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )

# def export_to_csv(df, filename="output.csv"):
#     csv = df.to_csv(index=False).encode('utf-8')
#     st.download_button(
#         label="Export to CSV",
#         data=csv,
#         file_name=filename,
#         mime="text/csv"
#     )

# st.title("✨ AI Content Creator Toolkit")

# # --- Session State Initialization ---
# # Content Generation
# if "cg_platform" not in st.session_state:
#     st.session_state["cg_platform"] = "Facebook"
# if "cg_prompt" not in st.session_state:
#     st.session_state["cg_prompt"] = ""
# if "cg_output" not in st.session_state:
#     st.session_state["cg_output"] = ""

# # Niche & Keyword Research
# if "nk_topic" not in st.session_state:
#     st.session_state["nk_topic"] = ""
# if "nk_output" not in st.session_state:
#     st.session_state["nk_output"] = ""

# # Meme Generation
# if "meme_template" not in st.session_state:
#     st.session_state["meme_template"] = ""
# if "meme_top" not in st.session_state:
#     st.session_state["meme_top"] = ""
# if "meme_bottom" not in st.session_state:
#     st.session_state["meme_bottom"] = ""
# if "meme_url" not in st.session_state:
#     st.session_state["meme_url"] = ""

# # YouTube SEO Metadata
# if "yt_title" not in st.session_state:
#     st.session_state["yt_title"] = ""
# if "yt_script" not in st.session_state:
#     st.session_state["yt_script"] = ""
# if "yt_output" not in st.session_state:
#     st.session_state["yt_output"] = ""

# # Caption-to-Image Prompt
# if "caption_input" not in st.session_state:
#     st.session_state["caption_input"] = ""
# if "caption_output" not in st.session_state:
#     st.session_state["caption_output"] = ""

# # --- Features ---
# if feature == "Content Generation":
#     st.header("📝 Content Generation for Social Media")
#     platforms = ["Facebook", "Instagram", "TikTok", "YouTube"]
#     st.session_state["cg_platform"] = st.selectbox("Select Platform", platforms, index=platforms.index(st.session_state["cg_platform"]))
#     st.session_state["cg_prompt"] = st.text_area("Enter your content prompt or idea", value=st.session_state["cg_prompt"])
#     if st.button("Generate Content"):
#         model = get_gemini_model()
#         if model and st.session_state["cg_prompt"]:
#             with st.spinner("Generating content..."):
#                 response = model.generate_content(
#                     f"Generate engaging, creative, and platform-optimized content for {st.session_state['cg_platform']} based on this prompt: {st.session_state['cg_prompt']}. "
#                     "Make it catchy, concise, and include a relevant emoji."
#                 )
#                 st.session_state["cg_output"] = response.text
#         elif not st.session_state["cg_prompt"]:
#             st.error("Please enter a prompt.")
#     if st.session_state["cg_output"]:
#         st.markdown(st.session_state["cg_output"])
#         st.text_area("Copy Output", st.session_state["cg_output"], height=200, label_visibility="collapsed")

# elif feature == "Niche & Keyword Research":
#     st.header("🔍 Niche and Keyword Research for SEO")
#     st.session_state["nk_topic"] = st.text_input("Enter a topic or niche idea", value=st.session_state["nk_topic"])
#     if st.button("Research Niches & Keywords"):
#         model = get_gemini_model()
#         if model and st.session_state["nk_topic"]:
#             with st.spinner("Researching..."):
#                 trends = get_trending_keywords(st.session_state["nk_topic"], gprop='')
#                 trends_str = ", ".join(trends) if trends else "No trending keywords found"
#                 prompt = (
#                     f"Conduct niche research and suggest SEO-optimized keywords for the topic: {st.session_state['nk_topic']}. "
#                     f"Here are some trending keywords: {trends_str}. "
#                     "Include 5-10 niche ideas, top keywords with search volume estimates (low/medium/high), "
#                     "long-tail keyword suggestions, and actionable tips. Present the results in a neat markdown table and use emojis."
#                 )
#                 response = model.generate_content(prompt)
#                 st.session_state["nk_output"] = response.text
#                 st.subheader("Research Results:")
#                 st.markdown(st.session_state["nk_output"])
#                 # Try to extract a markdown table for CSV export
#                 import re
#                 table_match = re.search(r"\|(.+\|)+\n(\|[-:]+)+\n((\|.*\|.*\n)+)", st.session_state["nk_output"])
#                 if table_match:
#                     table_md = table_match.group(0)
#                     try:
#                         df = pd.read_csv(pd.compat.StringIO(table_md.replace('|', ',')), skipinitialspace=True)
#                         export_to_csv(df, filename="niche_keywords.csv")
#                     except Exception:
#                         pass
#                 export_to_docx(st.session_state["nk_output"], filename="niche_keywords.docx")
#                 if trends:
#                     st.info(f"🔥 Trending keywords: {trends_str}")
#         elif not st.session_state["nk_topic"]:
#             st.error("Please enter a topic.")
#     if st.session_state["nk_output"]:
#         st.text_area("Copy Output", st.session_state["nk_output"], height=200, label_visibility="collapsed")

# elif feature == "Meme Generation":
#     st.header("😂 Meme Generation")
#     templates = get_meme_templates()
#     if templates:
#         template_names = [template['name'] for template in templates]
#         if not st.session_state["meme_template"]:
#             st.session_state["meme_template"] = template_names[0]
#         st.session_state["meme_template"] = st.selectbox("Select Meme Template", template_names, index=template_names.index(st.session_state["meme_template"]))
#         st.session_state["meme_top"] = st.text_input("Top Text (optional)", value=st.session_state["meme_top"])
#         st.session_state["meme_bottom"] = st.text_input("Bottom Text (optional)", value=st.session_state["meme_bottom"])
#         if st.button("Generate Meme"):
#             selected_template = st.session_state["meme_template"]
#             top_text = st.session_state["meme_top"]
#             bottom_text = st.session_state["meme_bottom"]
#             template_id = next((t['id'] for t in templates if t['name'] == selected_template), None)
#             if template_id:
#                 if not top_text and not bottom_text:
#                     model = get_gemini_model()
#                     if model:
#                         meme_trends = get_trending_keywords("meme", gprop='')
#                         meme_prompt = (
#                             f"Suggest a funny and relevant meme caption (top and bottom text) for the meme template '{selected_template}'. "
#                             f"Use one of these trending topics if possible: {', '.join(meme_trends)}. "
#                             "Return your answer as two lines: first line for top text, second line for bottom text. Avoid special characters."
#                         )
#                         meme_response = model.generate_content(meme_prompt)
#                         meme_lines = meme_response.text.strip().split('\n')
#                         top_text = meme_lines[0] if len(meme_lines) > 0 else "When you see AI"
#                         bottom_text = meme_lines[1] if len(meme_lines) > 1 else "And it actually works"
#                         st.session_state["meme_top"] = top_text
#                         st.session_state["meme_bottom"] = bottom_text
#                 elif not top_text:
#                     model = get_gemini_model()
#                     if model:
#                         meme_prompt = (
#                             f"Suggest a witty top meme caption for the template '{selected_template}' given the bottom text: '{bottom_text}'. "
#                             "Avoid special characters."
#                         )
#                         meme_response = model.generate_content(meme_prompt)
#                         top_text = meme_response.text.strip().split('\n')[0]
#                         st.session_state["meme_top"] = top_text
#                 elif not bottom_text:
#                     model = get_gemini_model()
#                     if model:
#                         meme_prompt = (
#                             f"Suggest a witty bottom meme caption for the template '{selected_template}' given the top text: '{top_text}'. "
#                             "Avoid special characters."
#                         )
#                         meme_response = model.generate_content(meme_prompt)
#                         bottom_text = meme_response.text.strip().split('\n')[0]
#                         st.session_state["meme_bottom"] = bottom_text

#                 def format_text(text):
#                     if not text:
#                         return "_"
#                     return (
#                         text.replace("-", "--")
#                             .replace("_", "__")
#                             .replace(" ", "_")
#                             .replace("?", "~q")
#                             .replace("%", "~p")
#                             .replace("#", "~h")
#                             .replace("/", "~s")
#                             .replace("\"", "''")
#                     )
#                 top_text_formatted = format_text(st.session_state["meme_top"])
#                 bottom_text_formatted = format_text(st.session_state["meme_bottom"])
#                 meme_url = f"https://api.memegen.link/images/{template_id}/{top_text_formatted}/{bottom_text_formatted}.png"
#                 st.session_state["meme_url"] = meme_url
#             else:
#                 st.error("Template ID not found.")
#         if st.session_state["meme_url"]:
#             st.subheader("Generated Meme:")
#             st.image(st.session_state["meme_url"], width='stretch')
#             st.markdown(f"[🔗 View Meme]({st.session_state['meme_url']})")

# elif feature == "YouTube SEO Metadata":
#     st.header("🎬 YouTube SEO Metadata Generator")
#     st.session_state["yt_title"] = st.text_input("Rough Video Title", value=st.session_state["yt_title"])
#     st.session_state["yt_script"] = st.text_area("Video Script or Description Outline", value=st.session_state["yt_script"])
#     if st.button("Generate SEO Metadata"):
#         model = get_gemini_model()
#         if model and st.session_state["yt_title"] and st.session_state["yt_script"]:
#             with st.spinner("Generating metadata..."):
#                 yt_trends = get_trending_keywords(st.session_state["yt_title"], gprop='youtube')
#                 yt_trends_str = ", ".join(yt_trends) if yt_trends else "No trending keywords found"
#                 prompt = (
#                     f"Generate SEO-optimized YouTube metadata for the following video:\n"
#                     f"Title: {st.session_state['yt_title']}\n"
#                     f"Script/Outline: {st.session_state['yt_script']}\n"
#                     f"Trending YouTube keywords: {yt_trends_str}\n"
#                     "Return the result in this markdown format with relevant emojis:\n\n"
#                     "### 📢 Optimized Title\n"
#                     "- (title here)\n\n"
#                     "### 📝 Description\n"
#                     "- (250-300 words, include trending keywords, calls to action, and use emojis)\n\n"
#                     "### 🏷️ Tags\n"
#                     "- tag1, tag2, tag3, ... (15-21 tags, comma separated)\n"
#                 )
#                 response = model.generate_content(prompt)
#                 st.session_state["yt_output"] = response.text
#                 st.markdown(st.session_state["yt_output"])
#                 if yt_trends:
#                     st.info(f"🔥 Trending YouTube keywords: {yt_trends_str}")
#         elif not st.session_state["yt_title"] or not st.session_state["yt_script"]:
#             st.error("Please provide both a rough title and video script.")
#     if st.session_state["yt_output"]:
#         st.text_area("Copy Output", st.session_state["yt_output"], height=200, label_visibility="collapsed")

# elif feature == "Caption-to-Image Prompt":
#     st.header("🎨 Caption-to-Image Prompt Generator")
#     st.session_state["caption_input"] = st.text_area("Enter your caption or idea for AI art (e.g., for MidJourney/Stable Diffusion)", value=st.session_state["caption_input"])
#     if st.button("Generate Image Prompt"):
#         model = get_gemini_model()
#         if model and st.session_state["caption_input"]:
#             with st.spinner("Generating image prompt..."):
#                 prompt = (
#                     f"Turn this caption into a detailed, creative AI art prompt suitable for MidJourney or Stable Diffusion. "
#                     f"Make it vivid, descriptive, and include style or mood if possible. Caption: {st.session_state['caption_input']}"
#                 )
#                 response = model.generate_content(prompt)
#                 st.session_state["caption_output"] = response.text
#                 st.markdown(st.session_state["caption_output"])
#         elif not st.session_state["caption_input"]:
#             st.error("Please enter a caption.")
#     if st.session_state["caption_output"]:
#         st.text_area("Copy Output", st.session_state["caption_output"], height=200, label_visibility="collapsed")


import streamlit as st
import google.generativeai as genai
import requests
import os

from dotenv import load_dotenv
from pytrends.request import TrendReq
import pandas as pd
from docx import Document
from io import BytesIO

load_dotenv()

st.set_page_config(
    page_title="AI Content Generator",
    page_icon="✨",
    layout="centered",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
        body, .stApp {
            background-color: #181824;
            color: #f1f1f1;
        }
        .stButton>button {
            background-color: #ff4b4b;
            color: white;
            border-radius: 8px;
            font-weight: bold;
        }
        .stTextInput>div>input, .stTextArea>div>textarea, .stSelectbox>div>div>div>div {
            background-color: #232339;
            color: #f1f1f1;
            border-radius: 6px;
        }
        .stSidebar {
            background-color: #232339;
        }
        .stRadio > div {
            gap: 0.5rem;
        }
        .stMarkdown a {
            color: #ff4b4b;
        }
    </style>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div style="text-align:center">
    <h2 style="color:#ff4b4b;">Creator Toolkit</h2>
    <p style="font-size:15px;">Your all-in-one AI-powered content lab.<br>Generate, research, and optimize for every platform!</p>
    <hr>
</div>
""", unsafe_allow_html=True)

st.sidebar.title("🛠️ Features")
feature = st.sidebar.radio(
    "Select a feature:",
    [
        "Content Generation",
        "Niche & Keyword Research",
        "Meme Generation",
        "YouTube SEO Metadata",
        "Caption-to-Image Prompt",
        "YouTube Thumbnail Generator"
    ]
)

gemini_api_key = os.getenv("GEMINI_API_KEY")
if not gemini_api_key:
    st.error("Please set your GEMINI_API_KEY environment variable to use AI features. For example, export GEMINI_API_KEY=your_key_here")

def get_gemini_model():
    if gemini_api_key:
        genai.configure(api_key=gemini_api_key)
        return genai.GenerativeModel('gemini-1.5-flash')
    return None

def get_trending_keywords(query, gprop='youtube'):
    try:
        pytrends = TrendReq(hl='en-US', tz=360)
        pytrends.build_payload([query], cat=0, timeframe='now 7-d', geo='', gprop=gprop)
        related = pytrends.related_queries()
        if query in related and related[query]['top'] is not None:
            return related[query]['top']['query'].tolist()[:10]
    except Exception:
        pass
    return []

@st.cache_data
def get_meme_templates():
    response = requests.get("https://api.memegen.link/templates/")
    if response.status_code == 200:
        return response.json()
    else:
        st.error("Failed to fetch meme templates from Memegen.link.")
        return []

def export_to_docx(text, filename="output.docx"):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="Export to DOCX",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def export_to_csv(df, filename="output.csv"):
    csv = df.to_csv(index=False).encode('utf-8')
    st.download_button(
        label="Export to CSV",
        data=csv,
        file_name=filename,
        mime="text/csv"
    )

def pollinations_image_url(prompt):
    # Pollinations API for thumbnail generation (SDXL)
    return f"https://image.pollinations.ai/prompt/{requests.utils.quote(prompt)}"

st.title("✨ AI Content Creator Toolkit")

# --- Session State Initialization ---
# Content Generation
if "cg_platform" not in st.session_state:
    st.session_state["cg_platform"] = "Facebook"
if "cg_prompt" not in st.session_state:
    st.session_state["cg_prompt"] = ""
if "cg_output" not in st.session_state:
    st.session_state["cg_output"] = ""

# Niche & Keyword Research
if "nk_topic" not in st.session_state:
    st.session_state["nk_topic"] = ""
if "nk_output" not in st.session_state:
    st.session_state["nk_output"] = ""

# Meme Generation
if "meme_template" not in st.session_state:
    st.session_state["meme_template"] = ""
if "meme_top" not in st.session_state:
    st.session_state["meme_top"] = ""
if "meme_bottom" not in st.session_state:
    st.session_state["meme_bottom"] = ""
if "meme_url" not in st.session_state:
    st.session_state["meme_url"] = ""

# YouTube SEO Metadata
if "yt_title" not in st.session_state:
    st.session_state["yt_title"] = ""
if "yt_script" not in st.session_state:
    st.session_state["yt_script"] = ""
if "yt_output" not in st.session_state:
    st.session_state["yt_output"] = ""
if "yt_thumbnail_prompt" not in st.session_state:
    st.session_state["yt_thumbnail_prompt"] = ""
if "yt_thumbnail_url" not in st.session_state:
    st.session_state["yt_thumbnail_url"] = ""

# Caption-to-Image Prompt
if "caption_input" not in st.session_state:
    st.session_state["caption_input"] = ""
if "caption_output" not in st.session_state:
    st.session_state["caption_output"] = ""

# --- Features ---
if feature == "Content Generation":
    st.header("📝 Content Generation for Social Media")
    platforms = ["Facebook", "Instagram", "TikTok", "YouTube"]
    st.session_state["cg_platform"] = st.selectbox("Select Platform", platforms, index=platforms.index(st.session_state["cg_platform"]))
    st.session_state["cg_prompt"] = st.text_area("Enter your content prompt or idea", value=st.session_state["cg_prompt"])
    if st.button("Generate Content"):
        model = get_gemini_model()
        if model and st.session_state["cg_prompt"]:
            with st.spinner("Generating content..."):
                response = model.generate_content(
                    f"Generate engaging, creative, and platform-optimized content for {st.session_state['cg_platform']} based on this prompt: {st.session_state['cg_prompt']}. "
                    "Make it catchy, concise, and include a relevant emoji."
                )
                st.session_state["cg_output"] = response.text
        elif not st.session_state["cg_prompt"]:
            st.error("Please enter a prompt.")
    if st.session_state["cg_output"]:
        st.markdown(st.session_state["cg_output"])
        st.text_area("Copy Output", st.session_state["cg_output"], height=200, label_visibility="collapsed")

elif feature == "Niche & Keyword Research":
    st.header("🔍 Niche and Keyword Research for SEO")
    st.session_state["nk_topic"] = st.text_input("Enter a topic or niche idea", value=st.session_state["nk_topic"])
    if st.button("Research Niches & Keywords"):
        model = get_gemini_model()
        if model and st.session_state["nk_topic"]:
            with st.spinner("Researching..."):
                trends = get_trending_keywords(st.session_state["nk_topic"], gprop='')
                trends_str = ", ".join(trends) if trends else "No trending keywords found"
                prompt = (
                    f"Conduct niche research and suggest SEO-optimized keywords for the topic: {st.session_state['nk_topic']}. "
                    f"Here are some trending keywords: {trends_str}. "
                    "Include 5-10 niche ideas, top keywords with search volume estimates (low/medium/high), "
                    "long-tail keyword suggestions, and actionable tips. Present the results in a neat markdown table and use emojis."
                )
                response = model.generate_content(prompt)
                st.session_state["nk_output"] = response.text
                st.subheader("Research Results:")
                st.markdown(st.session_state["nk_output"])
                # Try to extract a markdown table for CSV export
                import re
                table_match = re.search(r"\|(.+\|)+\n(\|[-:]+)+\n((\|.*\|.*\n)+)", st.session_state["nk_output"])
                if table_match:
                    table_md = table_match.group(0)
                    try:
                        df = pd.read_csv(pd.compat.StringIO(table_md.replace('|', ',')), skipinitialspace=True)
                        export_to_csv(df, filename="niche_keywords.csv")
                    except Exception:
                        pass
                export_to_docx(st.session_state["nk_output"], filename="niche_keywords.docx")
                if trends:
                    st.info(f"🔥 Trending keywords: {trends_str}")
        elif not st.session_state["nk_topic"]:
            st.error("Please enter a topic.")
    if st.session_state["nk_output"]:
        st.text_area("Copy Output", st.session_state["nk_output"], height=200, label_visibility="collapsed")

elif feature == "Meme Generation":
    st.header("😂 Meme Generation")
    templates = get_meme_templates()
    if templates:
        template_names = [template['name'] for template in templates]
        if not st.session_state["meme_template"]:
            st.session_state["meme_template"] = template_names[0]
        st.session_state["meme_template"] = st.selectbox("Select Meme Template", template_names, index=template_names.index(st.session_state["meme_template"]))
        st.session_state["meme_top"] = st.text_input("Top Text (optional)", value=st.session_state["meme_top"])
        st.session_state["meme_bottom"] = st.text_input("Bottom Text (optional)", value=st.session_state["meme_bottom"])
        if st.button("Generate Meme"):
            selected_template = st.session_state["meme_template"]
            top_text = st.session_state["meme_top"]
            bottom_text = st.session_state["meme_bottom"]
            template_id = next((t['id'] for t in templates if t['name'] == selected_template), None)
            if template_id:
                if not top_text and not bottom_text:
                    model = get_gemini_model()
                    if model:
                        meme_trends = get_trending_keywords("meme", gprop='')
                        meme_prompt = (
                            f"Suggest a funny and relevant meme caption (top and bottom text) for the meme template '{selected_template}'. "
                            f"Use one of these trending topics if possible: {', '.join(meme_trends)}. "
                            "Return your answer as two lines: first line for top text, second line for bottom text. Avoid special characters."
                        )
                        meme_response = model.generate_content(meme_prompt)
                        meme_lines = meme_response.text.strip().split('\n')
                        top_text = meme_lines[0] if len(meme_lines) > 0 else "When you see AI"
                        bottom_text = meme_lines[1] if len(meme_lines) > 1 else "And it actually works"
                        st.session_state["meme_top"] = top_text
                        st.session_state["meme_bottom"] = bottom_text
                elif not top_text:
                    model = get_gemini_model()
                    if model:
                        meme_prompt = (
                            f"Suggest a witty top meme caption for the template '{selected_template}' given the bottom text: '{bottom_text}'. "
                            "Avoid special characters."
                        )
                        meme_response = model.generate_content(meme_prompt)
                        top_text = meme_response.text.strip().split('\n')[0]
                        st.session_state["meme_top"] = top_text
                elif not bottom_text:
                    model = get_gemini_model()
                    if model:
                        meme_prompt = (
                            f"Suggest a witty bottom meme caption for the template '{selected_template}' given the top text: '{top_text}'. "
                            "Avoid special characters."
                        )
                        meme_response = model.generate_content(meme_prompt)
                        bottom_text = meme_response.text.strip().split('\n')[0]
                        st.session_state["meme_bottom"] = bottom_text

                def format_text(text):
                    if not text:
                        return "_"
                    return (
                        text.replace("-", "--")
                            .replace("_", "__")
                            .replace(" ", "_")
                            .replace("?", "~q")
                            .replace("%", "~p")
                            .replace("#", "~h")
                            .replace("/", "~s")
                            .replace("\"", "''")
                    )
                top_text_formatted = format_text(st.session_state["meme_top"])
                bottom_text_formatted = format_text(st.session_state["meme_bottom"])
                meme_url = f"https://api.memegen.link/images/{template_id}/{top_text_formatted}/{bottom_text_formatted}.png"
                st.session_state["meme_url"] = meme_url
            else:
                st.error("Template ID not found.")
        if st.session_state["meme_url"]:
            st.subheader("Generated Meme:")
            st.image(st.session_state["meme_url"], width='stretch')
            st.markdown(f"[🔗 View Meme]({st.session_state['meme_url']})")

elif feature == "YouTube SEO Metadata":
    st.header("🎬 YouTube SEO Metadata Generator")
    st.session_state["yt_title"] = st.text_input("Rough Video Title", value=st.session_state["yt_title"])
    st.session_state["yt_script"] = st.text_area("Video Script or Description Outline", value=st.session_state["yt_script"])
    if st.button("Generate SEO Metadata"):
        model = get_gemini_model()
        if model and st.session_state["yt_title"] and st.session_state["yt_script"]:
            with st.spinner("Generating metadata..."):
                yt_trends = get_trending_keywords(st.session_state["yt_title"], gprop='youtube')
                yt_trends_str = ", ".join(yt_trends) if yt_trends else "No trending keywords found"
                prompt = (
                    f"Generate SEO-optimized YouTube metadata for the following video:\n"
                    f"Title: {st.session_state['yt_title']}\n"
                    f"Script/Outline: {st.session_state['yt_script']}\n"
                    f"Trending YouTube keywords: {yt_trends_str}\n"
                    "Return the result in this markdown format with relevant emojis:\n\n"
                    "### 📢 Optimized Title\n"
                    "- (title here)\n\n"
                    "### 📝 Description\n"
                    "- (250-300 words, include trending keywords, calls to action, and use emojis)\n\n"
                    "### 🏷️ Tags\n"
                    "- tag1, tag2, tag3, ... (15-21 tags, comma separated)\n"
                )
                response = model.generate_content(prompt)
                st.session_state["yt_output"] = response.text
                st.markdown(st.session_state["yt_output"])
                if yt_trends:
                    st.info(f"🔥 Trending YouTube keywords: {yt_trends_str}")
    if st.session_state["yt_output"]:
        st.text_area("Copy Output", st.session_state["yt_output"], height=200, label_visibility="collapsed")
        # --- Thumbnail Generation Section ---
        st.markdown("---")
        st.subheader("🖼️ Generate YouTube Thumbnail from Metadata")
        if st.button("Generate Thumbnail Prompt & Image"):
            model = get_gemini_model()
            if model and st.session_state["yt_output"]:
                with st.spinner("Generating thumbnail prompt..."):
                    thumb_prompt = model.generate_content(
                        f"Based on this YouTube video metadata, write a detailed, creative prompt for an AI image generator to create a YouTube thumbnail. "
                        f"Focus on the main topic, mood, and visual elements. Metadata:\n{st.session_state['yt_output']}\n"
                        "The prompt should be suitable for Stable Diffusion XL or Pollinations API, and should be visually descriptive, concise, and avoid any text overlays."
                    ).text
                    st.session_state["yt_thumbnail_prompt"] = thumb_prompt
                    st.markdown("**Thumbnail Prompt:**")
                    st.markdown(thumb_prompt)
                    st.text_area("Copy Thumbnail Prompt", thumb_prompt, height=120, label_visibility="collapsed")
                    # Generate thumbnail image using Pollinations API
                    st.session_state["yt_thumbnail_url"] = pollinations_image_url(thumb_prompt)
                    st.markdown("**Generated Thumbnail:**")
                    st.image(st.session_state["yt_thumbnail_url"], caption="AI-Generated Thumbnail", width='stretch')
                    st.markdown(f"[🔗 View Full Image]({st.session_state['yt_thumbnail_url']})")
        # Show last generated thumbnail if available
        if st.session_state["yt_thumbnail_prompt"]:
            st.markdown("**Last Generated Thumbnail Prompt:**")
            st.markdown(st.session_state["yt_thumbnail_prompt"])
            st.text_area("Copy Thumbnail Prompt", st.session_state["yt_thumbnail_prompt"], height=120, label_visibility="collapsed")
        if st.session_state["yt_thumbnail_url"]:
            st.markdown("**Last Generated Thumbnail:**")
            st.image(st.session_state["yt_thumbnail_url"], caption="AI-Generated Thumbnail", width='stretch')
            st.markdown(f"[🔗 View Full Image]({st.session_state['yt_thumbnail_url']})")

elif feature == "Caption-to-Image Prompt":
    st.header("🎨 Caption-to-Image Prompt Generator")
    st.session_state["caption_input"] = st.text_area("Enter your caption or idea for AI art (e.g., for MidJourney/Stable Diffusion)", value=st.session_state["caption_input"])
    if st.button("Generate Image Prompt"):
        model = get_gemini_model()
        if model and st.session_state["caption_input"]:
            with st.spinner("Generating image prompt..."):
                prompt = (
                    f"Turn this caption into a detailed, creative AI art prompt suitable for MidJourney or Stable Diffusion. "
                    f"Make it vivid, descriptive, and include style or mood if possible. Caption: {st.session_state['caption_input']}"
                )
                response = model.generate_content(prompt)
                st.session_state["caption_output"] = response.text
                st.markdown(st.session_state["caption_output"])
    if st.session_state["caption_output"]:
        st.text_area("Copy Output", st.session_state["caption_output"], height=200, label_visibility="collapsed")

# --- YouTube Thumbnail Generator Standalone Feature (optional direct access) ---
if feature == "YouTube Thumbnail Generator":
    st.header("🖼️ YouTube Thumbnail Generator")
    st.info("Generate a thumbnail prompt and image for your YouTube video using Gemini and Pollinations API.")
    yt_metadata = st.text_area("Paste your YouTube metadata (title, description, tags, etc.) here")
    if st.button("Generate Thumbnail from Metadata"):
        model = get_gemini_model()
        if model and yt_metadata:
            with st.spinner("Generating thumbnail prompt..."):
                thumb_prompt = model.generate_content(
                f"""
                # ROLE: You are a YouTube Thumbnail Artist specializing in creating viral, high-click-through-rate (CTR) thumbnails using AI image generators.

                # TASK: Analyze the provided YouTube video metadata and generate ONE single, highly descriptive image prompt. The prompt must visualize the core concept of the video as a dramatic, attention-grabbing contrast.

                # METADATA:
                {yt_metadata}

                # INSTRUCTIONS:
                1.  Identify the main subject (e.g., a person, product, concept).
                2.  Identify the core conflict or comparison (e.g., "X vs Y", "good vs bad", "before vs after", "shocking result").
                3.  Translate this conflict into a strong visual metaphor using a split-screen, a glowing vs dark object, a transformation, or a shocking facial expression.
                4.  Inject high-energy keywords to maximize visual appeal.
                5.  **OUTPUT ONLY THE PROMPT. NO OTHER TEXT.**

                # PROMPT TEMPLATE:
                "[Genre] of [Main Subject], [Visualizing the Conflict]. [Style and Quality Boosters]."

                # EXAMPLE 1 (For an 'AI Tools for Creators' video):
                "YouTube thumbnail of a man's face split into four quadrants, each with a different visual style and a glowing AI logo (VidIQ, Canva, ChatGPT, Copilot), one quadrant is photorealistic and clear while the others are distorted and glitched. Cinematic lighting, hyper-detailed, digital art, trending on ArtStation."

                # EXAMPLE 2 (For a 'Python vs Javascript' video):
                "YouTube thumbnail of a fierce visual battle between a sleek Python snake and a detailed Javascript shield, cracking on impact. Epic dynamic lighting, sparks flying, hyper-realistic, 3D render, dramatic atmosphere."

                # EXAMPLE 3 (For a 'We Tested 5 Gaming Chairs' video):
                "YouTube thumbnail of a single gamer reacting with a shocked expression, sitting in five different gaming chairs that are visually merging together around him. Vibrant colors, depth of field, studio lighting, photorealistic."

        #YOUR OUTPUT (Follow the template and examples. Be creative and dramatic):
                """
                ).text


                st.markdown("**Thumbnail Prompt:**")
                st.markdown(thumb_prompt)
                # st.text_area("Copy Thumbnail Prompt", thumb_prompt, height=120, label_visibility="collapsed")
                st.text_area("Copy Thumbnail Prompt", st.session_state["yt_thumbnail_prompt"], height=120, label_visibility="collapsed", key="yt_thumb_copy")
                # Generate thumbnail image using Pollinations API
                thumb_url = pollinations_image_url(thumb_prompt)
                st.markdown("**Generated Thumbnail:**")
                st.image(thumb_url, caption="AI-Generated Thumbnail", width='stretch')
                st.markdown(f"[🔗 View Full Image]({thumb_url})")